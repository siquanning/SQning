from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("SST_上机测试步骤.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
RED = "9B1C1C"
GOLD = "7A5A00"
BLACK = "202020"
MUTED = "606A73"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=10, bold=False, color=BLACK, name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, bold_prefix=None, color=BLACK, after=3, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.keep_with_next = keep
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True, color=color)
        set_font(p.add_run(text[len(bold_prefix):]), color=color)
    else:
        set_font(p.add_run(text), color=color)
    return p


def add_bullets(doc, items):
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08
        set_font(p.add_run(text))


def add_steps(doc, items):
    for text in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
        set_font(p.add_run(text))


def add_code(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.0
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    set_font(p.add_run("\n".join(lines)), size=8.5, name="Consolas", color="30343B")


def add_callout(doc, title, text, risk=False):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FDECEC" if risk else "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(title + "  "), bold=True, color=RED if risk else DARK_BLUE)
    set_font(p.add_run(text), size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 15, 7),
        ("Heading 2", 12.5, BLUE, 10, 5),
        ("Heading 3", 11, DARK_BLUE, 7, 3),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(3)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_result_table(doc):
    rows = [
        ("Debug全量构建与下载", "0 errors，确认下载最新Debug输出"),
        ("上电安全初态", "GPIO20/22/23=0，PWM Block"),
        ("12路ADC offset", "零输入换算接近0，各通道独立"),
        ("Vdc比例", "多点线性及CT1/CT2/Gain正确"),
        ("PLL", "50Hz、vq小、req=1、alpha≈1"),
        ("PRECHARGE", "六路上升，最低值达到门槛"),
        ("BYPASS_WAIT", "GPIO23=1，延时期间PWM Block"),
        ("RUN", "PLL/TZ/FAULT许可后才释放PWM"),
        ("STOP三阶段", "PWM→GPIO23→GPIO22顺序正确"),
        ("TZ故障", "硬件封锁、FAULT锁存、不自启"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [2300, 4200, 1200, 1660])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(("阶段", "关键验收结果", "结论", "记录/截图号")):
        set_cell_shading(hdr.cells[i], LIGHT_BLUE)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(text), size=9, bold=True, color=DARK_BLUE)
    for stage, criterion in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((stage, criterion, "□通过  □失败", "")):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(text), size=8.5)
    set_table_geometry(table, [2300, 4200, 1200, 1660])


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.68)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)
configure_styles(doc)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("SST 前级整流控制 · 上机测试规程"), size=8.5, color=MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
footer._p.append(field)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(2)
set_font(p.add_run("TMS320F28335 SST"), size=11, bold=True, color=BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
set_font(p.add_run("前级整流控制上机测试步骤"), size=24, bold=True, color=DARK_BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
set_font(p.add_run("工程：E:\\repos\\resst　目标板：TMS320F28335　版本：低压分阶段验证版"), size=9, color=MUTED)

add_callout(doc, "安全原则", "低压、限流、空载、逐级放行。未经前一级验证通过不得进入下一级；功率接线、放电和高压测量必须由具备资质人员完成。", risk=True)

add_heading(doc, "1　测试目标与禁止事项", 1)
add_text(doc, "依次验证固件与安全初态、12路ADC零偏、Vdc测量比例、PLL锁相、GPIO22预充、GPIO23旁路、PWM释放、GPIO21停止及TZ故障保护。")
add_bullets(doc, [
    "GPIO21停止后必须先PWM_BlockOutput()，再断GPIO23，最后断GPIO22。",
    "TZ1/TZ2异常时禁止释放PWM；预充、旁路等待和PLL等待期间PWM必须保持Block。",
    "不得用ADC offset补偿CT1/CT2/Gain比例误差；未确认真实Vdc比例前不得把门槛直接填成20V、100V或800V。",
    "FAULT不得自动清除；先断功率、确认放电、查明原因，再复位DSP。",
])

add_heading(doc, "2　设备、接线与上电前检查", 1)
add_bullets(doc, [
    "隔离可调且带限流的低压三相电源；示波器、隔离差分探头、电流探头、万用表；XDS、CCS和串口工具。",
    "GPIO21：保持型高有效启动；GPIO20：运行指示；GPIO22：S1/S2/S3输入总开关；GPIO23：S4/S5/S6预充旁路总开关。",
    "首次测试断开功率输入，确认6路直流电容已放电；仅保持DSP控制板和采样板供电。",
])

add_heading(doc, "3　编译、下载与安全初态", 1)
add_steps(doc, [
    "CCS选择Debug并执行Full Build，要求0 errors；下载最新Debug\\resst.out，不使用旧Flash_Release输出。",
    "GPIO21保持0后复位并Resume，等待至少100ms，使初始化及50ms低电平消抖完成。",
    "确认GPIO20=0、GPIO22=0、GPIO23=0；示波器确认所有ePWM门极输出被封锁。",
    "核对系统状态、ADC帧、20kHz ISR及调度诊断。",
])
add_code(doc, [
    "Diagnostics_Get()->system_state    = 2 (STANDBY)",
    "Diagnostics_Get()->fault_code      = 0 (FAULT_NONE)",
    "fast_isr_count / g_adc_frame_count 持续增加",
    "g_start_seq_state=0, g_grid_switch_cmd=0, g_bypass_switch_cmd=0",
    "fast_isr.max_cycles 明显低于7000；miss计数不持续增加",
])

add_heading(doc, "4　12路ADC零偏整定", 1)
add_heading(doc, "4.1　零输入与Modbus采样", 2)
add_text(doc, "确认功率输入断开、母线放电，Vac=0、Iac=0。不能证明输入为0时，禁止校准该通道。若使用Modbus，将BOARD_DEBUG_JUSTFLOAT_ENABLE设为0后重新完整编译。")
add_code(doc, [
    "发送：01 03 00 00 00 0C 45 CF",
    "HR0~HR11：Vdc1~Vdc6、Vab/Vbc/Vca、Ia~Ic；16位数据高字节在前",
    "连续采集至少100帧，各通道独立求平均并四舍五入",
])
add_heading(doc, "4.2　CCS写入与验证", 2)
add_code(doc, [
    "g_vdc1_offset_counts ... g_vdc6_offset_counts",
    "g_vac_vab_offset_counts / g_vac_vbc_offset_counts / g_vac_vca_offset_counts",
    "g_iac_ia_offset_counts / g_iac_ib_offset_counts / g_iac_ic_offset_counts",
])
add_text(doc, "单位均为ADC count。Vac零偏只校正线电压三路ADC（Vab/Vbc/Vca），看g_measurement.vline_v或JustFloat lite mode 0；禁止对着重构相电压vac_v或mode 2调零偏。写入raw平均值后，Vdc应接近0且不为负，vline/Iac应围绕0小幅波动。确认后写回board_config.h对应*_OFFSET_COUNTS_DEFAULT，再编译、复位并复查。")
add_callout(doc, "偏置纪律", "Offset只修正零点。若输入翻倍而软件结果不近似翻倍，应检查CT1/CT2/Gain，禁止继续调offset。", risk=True)

add_heading(doc, "5　Vdc、Vac、Iac比例标定与预充参数整定", 1)
add_text(doc, "零偏整定完成后，必须分别验证Vdc、Vac和Iac三类通道的比例及极性。三类量共用ADC Vref/最大码值，但各自使用独立的CT1、CT2、模拟调理增益和换算公式；不得只校Vdc后假定Vac/Iac也正确。")
add_heading(doc, "5.1　当前工程标定值与公式", 2)
add_text(doc, "Vdc当前实装标定：CT1=1000V:2V、CT2=1:1、实测后级模拟Gain=0.5，因此软件值为按实测raw换算的一次侧电压，理论满量程约3000V。")
add_code(doc, [
    "Vdc = max(raw - offset, 0) × 3 / 4095 × 1000 / 2",
    "    ≈ corrected_raw × 0.73260 V",
    "400V门槛约对应 corrected_raw = 546 counts",
])
add_text(doc, "Vac当前配置：CT1=3000Vrms:100Vrms；CT2=100Vrms:0.00167Arms；TIA=1000Ω；模拟Gain=0.5。换算为有符号瞬时一次侧电压：")
add_code(doc, [
    "Vac = polarity × (raw - offset) × VREF/MAX_COUNT",
    "      ÷ (VAC_TIA_OHM × VAC_ANALOG_GAIN)",
    "      × (VAC_CT2_PRI_V_RMS/VAC_CT2_SEC_A_RMS)",
    "      × (VAC_CT1_PRI_V_RMS/VAC_CT1_SEC_V_RMS)",
    "当前约为 polarity × delta × 2.6322 V/count",
])
add_text(doc, "Iac当前配置：CT1=100A:5A；CT2=5A:0.0025A；TIA=1000Ω；模拟Gain=0.5。换算为有符号一次侧电流：")
add_code(doc, [
    "Iac = polarity × (raw - offset) × VREF/MAX_COUNT",
    "      ÷ (IAC_TIA_OHM × IAC_ANALOG_GAIN)",
    "      × (IAC_CT2_PRI_A/IAC_CT2_SEC_A)",
    "      × (IAC_CT1_PRI_A/IAC_CT1_SEC_A)",
    "当前约为 polarity × delta × 0.058608 A/count",
])
add_heading(doc, "5.2　三类通道实测标定", 2)
add_steps(doc, [
    "Vdc：逐路施加两个以上已知直流低压点（如0.5V、1.0V、1.5V），记录raw、g_measurement.vdc_v[]及万用表值；不得超过ADC 0~3V。按真实硬件修改BOARD_VDC_CT1_PRI_V/SEC_V、BOARD_VDC_CT2_PRI_V/SEC_V、BOARD_VDC_ANALOG_GAIN。",
    "Vac：使用隔离低压正弦源验证线电压，记录示波器/万用表有效值、raw波形及g_measurement.vline_v[]（Vab/Vbc/Vca）；零点必须在vline上确认。相序与PLL看重构后的g_measurement.vac_v[]。比例由BOARD_VAC_CT1_*、BOARD_VAC_CT2_*、BOARD_VAC_TIA_OHM、BOARD_VAC_ANALOG_GAIN控制。",
    "Iac：使用隔离限流回路和电流探头/标准表逐相注入至少两个电流点，记录raw及g_measurement.iac_a[]；确认幅值、正负方向和三相通道映射。比例由BOARD_IAC_CT1_*、BOARD_IAC_CT2_*、BOARD_IAC_TIA_OHM、BOARD_IAC_ANALOG_GAIN控制。",
    "BOARD_ADC_VREF_V和BOARD_ADC_MAX_COUNT影响三类测量。Vref应实测或依据硬件确认，不得为凑数随意修改。任何比例参数修改后，必须重新完整编译并复测三类通道。",
    "测量比例合格后，现场通过CCS调整g_precharge_done_v、g_precharge_timeout_ms、g_bypass_delay_ms；参数基本确定后再写回对应*_DEFAULT宏。",
])
add_callout(doc, "门槛设置", "当前BOARD_PRECHARGE_DONE_V_DEFAULT为400V，约对应1092个有效counts；含义是六路最低一路也必须达到400V，并非六路合计400V。低压实验前必须在CCS修改g_precharge_done_v，DSP复位后恢复400V。")

add_heading(doc, "6　PLL独立测试", 1)
add_text(doc, "GPIO21保持0，PWM Block，GPIO22/23保持0。向已确认安全的Vac采样输入位置注入隔离、限流、相序正确的三相50Hz信号。当前BOARD_PLL_LOCK_VMAG_MIN_V=50.0V，表示Clarke变换后的αβ矢量峰值；测试相电压至少约35.4Vrms（线电压约61.2Vrms）才能越过该门槛。不得把50V直接加到ADC引脚，也不得绕过PLL判据。")
add_code(doc, [
    "观察：g_measurement.vline_v[0..2]、g_measurement.vac_v[0..2]、g_pll.vmag、g_pll.freq",
    "      g_pll.vd、g_pll.vq、g_pll.theta、g_pll_switch_req",
    "      g_switch_alpha、g_switch_phase_err_deg",
])
add_bullets(doc, [
    "g_pll.freq≈50Hz不能单独证明锁定：vmag低于门槛时PLL会进入50Hz holdover。",
    "通过：三相幅值和相序正确；vmag>50V；频率稳定在50Hz附近；|vq|<3%×vmag且vd>90%×vmag。",
    "连续满足判据约200ms后g_pll_switch_req=1；随后约200ms内g_switch_alpha平滑升至1附近，无突跳和FAULT。",
    "撤去或降低Vac后按现有失锁/holdover逻辑变化，且不得发生非法PWM释放。",
    "不锁时依次查Vac offset、通道映射、相序、幅值门槛、频率、vd/vq和TZ；不要先改PI、theta、90°补偿或alpha算法。",
])

add_heading(doc, "7　不控整流软启动测试", 1)
add_heading(doc, "7.1　PRECHARGE", 2)
add_steps(doc, [
    "保持低压限流，且PLL独立测试已通过；GPIO21先保持0至少50ms。",
    "CCS加入预充参数、六路Vdc、最低值、启动状态/计时/失败标志及GPIO命令变量。",
    "GPIO21置1并保持：期望state=PRECHARGE，GPIO22=1、GPIO23=0、GPIO20=0，PWM仍Block。",
    "观察六路电容经预充电阻上升；g_precharge_vdc_min必须等于六路最低值。",
    "最低一路达到g_precharge_done_v后，GPIO23=1并进入BYPASS_WAIT。若超时，失败标志置位，GPIO23/22断开且禁止RUN；先松开GPIO21再排查。",
])
add_code(doc, [
    "g_precharge_done_v / g_precharge_timeout_ms / g_bypass_delay_ms",
    "g_precharge_vdc_min / g_start_seq_state / g_start_seq_timer_ms",
    "g_start_seq_fail / g_grid_switch_cmd / g_bypass_switch_cmd",
    "g_measurement.vdc_v[0] ... g_measurement.vdc_v[5]",
])
add_heading(doc, "7.2　BYPASS_WAIT与RUN", 2)
add_text(doc, "GPIO23闭合后，PWM仍必须Block至少g_bypass_delay_ms。延时结束后仅当PLL就绪、g_switch_alpha≥g_pll_ready_alpha_min、TZ1/TZ2正常且无FAULT，才可RequestRun并释放PWM。")
add_bullets(doc, [
    "延时期间门极无PWM；GPIO22/23保持1。",
    "全部许可满足后GPIO20=1、系统进入RUN、PWM波形出现。",
    "首次释放使用最低可行母线电压并严格限流；六路Vdc不得异常跌落，电流或波形异常立即STOP/急停。",
])

add_heading(doc, "8　STOP与TZ安全回归", 1)
add_text(doc, "分别在PRECHARGE、BYPASS_WAIT和RUN阶段将GPIO21置0并保持超过50ms，使用示波器验证以下固定顺序：")
add_code(doc, [
    "1. PWM立即Block，停止主动开关",
    "2. GPIO23=0，退出预充电阻旁路",
    "3. GPIO22=0，切断三相输入",
    "4. 启动状态清零，回STANDBY，GPIO20=0",
])
add_callout(doc, "安全顺序不可调整", "禁止GPIO22/23先断、PWM后封锁；三个阶段必须分别通过。", risk=True)
add_text(doc, "RUN阶段再进行受控TZ注入：TZ1或TZ2有效后，GPIO30应立即拉低，PWM由硬件OST封锁，系统进入FAULT，GPIO20熄灭且不得自动重启。故障后断功率、放电并复位恢复。")

add_heading(doc, "9　测试记录与最终放行", 1)
add_text(doc, "每次至少记录日期、测试人员、.out时间/CRC、供电与限流、12路offset、CT1/CT2/Gain、预充参数、PLL关键量、六路Vdc、ISR最大cycles、状态/故障码、GPIO时序及示波器截图。")
add_result_table(doc)
add_callout(doc, "最终放行", "全部项目通过且低压重复启停至少5次无异常，才可制定逐级升压计划。每一级均须重新确认Vdc比例、预充时间、电流峰值、旁路瞬态、PLL和TZ保护；不得从低压直接跳到高压满功率。", risk=True)

doc.core_properties.title = "TMS320F28335 SST 前级整流控制上机测试步骤"
doc.core_properties.subject = "参数整定、PLL、软启动及安全回归测试规程"
doc.core_properties.author = "SST项目组"
doc.save(OUT)
print(OUT)
